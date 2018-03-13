#!/usr/bin/env python3
# -*- coding:utf-8 -*- 
# Created by yaochao at 2018/3/8


import docx
import os
import re
from execjs_demo.execjs_demo import get_diff
import pprint

base = '/Users/yaochao/work/说明书差异对比和提取'
file1 = os.path.join(base, 'doc_handle1/【JYY编号1】硝苯地平控释片说明书 OCR版-质控后.docx')
file2 = os.path.join(base, 'doc_handle2/【JYY编号1】1.钙拮抗剂-1：[拜新同]硝苯地平控释片（done） 3.docx')
files_name = os.listdir(base)

# 所有文件title的映射表
titles_map = [['核准日期'], ['修改日期'], ['药品名称'], ['成份', '成 份'], ['适应症'], ['性状', '性 状'], ['规格', '规 格'], ['用法用量', '用法和用量'],
              ['按照患者体重50 kg计算'], ['不良反应'], ['禁忌'], ['注意事项', '注意事 项'],
              ['孕妇及哺乳期妇女用药', '孕妇与哺乳期妇女用药', '孕妇及晡乳期妇女用药', '孕妇和哺乳期妇女用药'], ['儿童用药'], ['老年用药', '老年患者用药'], ['药物相互作用'],
              ['药物过量'], ['临床试验'], ['药理毒理', '药物毒理'], ['药代动力学'], ['贮藏', '贮 藏', '贮   藏'], ['包装', '包 装'], ['有效期'], ['执行标准'],
              ['批准文号'], ['生产企业'], ['进口分装企业'], ['进口药品注册证号'], ['热线']]


def get_all_texts():
    files = [os.path.join(base, x) for x in files_name]
    texts = []
    for f in files:
        if '.docx' in f:
            doc = docx.Document(f)
            paragraphs = doc.paragraphs
            text = '\n'.join([x.text for x in paragraphs])
            texts.append(text)
    return texts


def get_text(f):
    doc = docx.Document(f)
    paragraphs = doc.paragraphs
    text = '\n'.join([x.text for x in paragraphs])
    return text


def get_all_titles(texts):
    title_re = r'【.+?】'
    title_re = re.compile(title_re)
    titles = []
    for t in texts:
        r = re.findall(title_re, t)
        titles.extend(r)
    return set(titles)


def get_all_parts(f):
    text = get_text(f)
    title_re = '【(.+?)】([\w\W]+?)(?<![见和])(?=【|$)'
    parts = re.findall(title_re, text)
    return parts


def title_unification(title):
    '''
    把给定的title归一为统一的title
    '''
    for i in titles_map:
        if title in i:
            return i[0]
    return title


def compare_two_files(f1, f2):
    # 1. 分别拆成不同的部分，归一标题
    parts1 = get_all_parts(f1)
    parts1 = [[title_unification(x[0]), x[1]] for x in parts1]
    parts1_titles = [x[0] for x in parts1]
    parts2 = get_all_parts(f2)
    parts2 = [[title_unification(x[0]), x[1]] for x in parts2]
    parts2_titles = [x[0] for x in parts2]
    # 2. for循环，对两个文本做合并，1为f1单独有，2为f2单独有，0为共有
    for idx, i in enumerate(parts1_titles):
        if i in parts2_titles:
            parts1_i = parts1[idx]
            parts1.pop(idx)
            # 把f1和f2中共有的title所对应的part，从f2中移除，方面循环parts2_titles时，剩下的全部是f2中特有的。
            parts2_i_index = parts2_titles.index(i)
            parts2_i = parts2[parts2_i_index]
            parts2_titles.pop(parts2_i_index)
            parts2.pop(parts2_i_index)
            # 组成新的parts1_i，即parts1中第i个元素，第一个是0，代表共有，第二个是parts1中的，第三个是parts2中的，重组后的形式 [0, '生产企业', '\n企业名称：Bayer Pharma AG \n生产地址：51368 Leverkusen,Germany \n电话号码：+49(0)214 30/57430 \n传真号码：+49(9)214 9657430\n', '拜耳医药保健有限公司\n']
            parts1_i = [0, parts1_i[0], parts1_i[1], parts2_i[1]]
            parts1.insert(idx, parts1_i)
        else:
            parts1[idx].insert(0, 1)
    for idx, i in enumerate(parts2_titles):
        parts2[idx].insert(0, 2)
        parts1.insert(0, parts2[idx])

    # 3. for 循环，对比文本的不同
    for i in parts1:
        if i[0] == 0:
            title = i[1]
            text1 = i[2].strip()
            text2 = i[3].strip()
            diff = get_diff(text1, text2)
            i.append(diff)
    return parts1


def main():
    parts = get_all_parts()
    for part in parts:
        print(part[0])


if __name__ == '__main__':
    compare_two_files(file1, file2)
